"""Utility to convert Markdown deliverables into DOCX files."""

from __future__ import annotations

from argparse import ArgumentParser
from pathlib import Path
from typing import Iterable, List

from docx import Document
from docx.shared import Pt

REPO_ROOT = Path(__file__).resolve().parents[1]
DEFAULT_FILES = [
    REPO_ROOT / "docs/entregables_dic_2025/01_analisis_flujo_musems.md",
    REPO_ROOT / "docs/entregables_dic_2025/02_plan_trabajo_pruebas_musems.md",
    REPO_ROOT / "docs/entregables_dic_2025/03_matriz_validaciones_musems.md",
    REPO_ROOT / "docs/entregables_dic_2025/04_scripts_validaciones_musems.md",
    REPO_ROOT / "docs/entregables_dic_2025/05_evidencia_ejecucion_scripts_musems.md",
    REPO_ROOT / "docs/entregables_dic_2025/06_resultados_validaciones_musems.md",
    REPO_ROOT / "docs/entregables_dic_2025/07_reporte_sistema_vida_saludable.md",
]
ALLOWED_TABLE_CHARS = set("|-: ")


def is_table_separator(line: str) -> bool:
    stripped = line.strip()
    if not stripped or "|" not in stripped:
        return False
    return set(stripped) <= ALLOWED_TABLE_CHARS


def parse_table(lines: List[str], start_idx: int) -> tuple[list[list[str]], int]:
    rows: list[list[str]] = []
    idx = start_idx
    while idx < len(lines):
        raw = lines[idx]
        if "|" not in raw:
            break
        rows.append([cell.strip() for cell in raw.strip().strip("|").split("|")])
        idx += 1
    return rows, idx


def write_table(document: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    table = document.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    for row_idx, row in enumerate(rows):
        for col_idx, cell in enumerate(row):
            paragraph = table.cell(row_idx, col_idx).paragraphs[0]
            paragraph.text = ""
            run = paragraph.add_run(cell)
            if row_idx == 0:
                run.bold = True


def add_code_paragraph(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    run = paragraph.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(10)


def convert_markdown(lines: List[str], document: Document) -> None:
    idx = 0
    in_code = False
    while idx < len(lines):
        raw = lines[idx]
        stripped = raw.strip()
        if stripped.startswith("```"):
            in_code = not in_code
            idx += 1
            continue
        if in_code:
            add_code_paragraph(document, raw)
            idx += 1
            continue
        if not stripped:
            document.add_paragraph("")
            idx += 1
            continue
        if stripped.startswith("#"):
            level = len(stripped) - len(stripped.lstrip("#"))
            content = stripped[level:].strip() or " "
            document.add_heading(content, level=min(level, 4))
            idx += 1
            continue
        if "|" in stripped and idx + 1 < len(lines) and is_table_separator(lines[idx + 1]):
            rows, next_idx = parse_table(lines, idx)
            write_table(document, rows)
            idx = next_idx
            continue
        if stripped.startswith(("- ", "* ")):
            document.add_paragraph(stripped[2:].strip(), style="List Bullet")
            idx += 1
            continue
        first_token = stripped.split()[0]
        if first_token.rstrip(".").isdigit():
            remainder = stripped[len(first_token):].strip()
            document.add_paragraph(remainder, style="List Number")
            idx += 1
            continue
        document.add_paragraph(raw)
        idx += 1


def convert_file(md_path: Path, output_path: Path | None = None) -> Path:
    lines = md_path.read_text(encoding="utf-8").splitlines()
    document = Document()
    convert_markdown(lines, document)
    target = output_path or md_path.with_suffix(".docx")
    document.save(target)
    return target


def resolve_targets(inputs: Iterable[str]) -> list[Path]:
    resolved: list[Path] = []
    for entry in inputs:
        path = Path(entry)
        if not path.is_absolute():
            path = REPO_ROOT / path
        if path.is_dir():
            resolved.extend(sorted(path.rglob("*.md")))
        elif path.suffix.lower() == ".md":
            resolved.append(path)
    return resolved


def main(argv: Iterable[str] | None = None) -> None:
    parser = ArgumentParser(description="Convert Markdown files to DOCX with table support.")
    parser.add_argument(
        "paths",
        nargs="*",
        help="Markdown files or directories to convert. Defaults to deliverables 01-06.",
    )
    args = parser.parse_args(list(argv) if argv is not None else None)
    targets = resolve_targets(args.paths or DEFAULT_FILES)
    if not targets:
        raise SystemExit("No Markdown files found to convert.")
    for file_path in targets:
        output = convert_file(file_path)
        rel = output.relative_to(REPO_ROOT)
        print(f"Converted {rel}")


if __name__ == "__main__":
    main()
